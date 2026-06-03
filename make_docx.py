
import sys, json, base64, io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

LOGO_SMALL = "iVBORw0KGgoAAAANSUhEUgAAAFAAAABQCAIAAAABc2X6AAAKHmlDQ1BJQ0MgUHJvZmlsZQAAeJy1Vnk8lGsbft73nX2xzZDd2LdGljDIvpPITpsxMxjLYMyg0iapcCJJthI5FTp0WpDTIi3ajtKmos7IEarT0SKVyvcOf+j7fefP812/3/O813v97vt+7ud+/3gvAMhjAAWMrhSBSBjs7caIjIpm4B8DBKgBRaAHtNicjDTwv4Dm6ceHc2/3mNLd+JPjs9Z3YS3Zbl/+vLHVjvoPuT9CjsvL4KDlPFC+NhY9HOVdKKfHhga7o/w+AAQKN4XLBYAoQfUd8bMxpARpTPwPMcniFD6q50j1FB47A+UlKNeLTUoTofyUVBfO5V6b5T/kingctB5pENUpmWIeehZJOpftWSJpLll6fzonTSjleSi35SSw0RjyWZQvnOt/FloZ0gH6errbWNjZ2DAtmRaM2GQ2J4mRwWEnS6v+25B+qzmmdxAAWbS3ttscsTBzTsNINywgAVlABypAE+gCI8AElsAWOAAX4An8QCAIBVFgNeCABJAChCAL5IAtIB8UghKwF1SBWtAAGkELOAHawVlwEVwFN8Ed8AAMAAkYAa/ABPgIpiEIwkNUiAapQFqQPmQKWUIsyAnyhJZCwVAUFAPFQwJIDOVAW6FCqBSqguqgRuhX6Ax0EboO9UGPoSFoHHoHfYERmALTYQ3YAF4Es2BX2B8OhVfB8XA6vA7Og3fBFXA9fAxugy/CN+EHsAR+BU8iACEjSog2wkRYiDsSiEQjcYgQ2YgUIOVIPdKCdCI9yD1EgrxGPmNwGBqGgWFiHDA+mDAMB5OO2YgpwlRhjmLaMJcx9zBDmAnMdywVq441xdpjfbGR2HhsFjYfW449jD2NvYJ9gB3BfsThcEo4Q5wtzgcXhUvErccV4fbjWnFduD7cMG4Sj8er4E3xjvhAPBsvwufjK/HH8Bfwd/Ej+E8EMkGLYEnwIkQTBIRcQjmhiXCecJcwSpgmyhH1ifbEQCKXuJZYTGwgdhJvE0eI0yR5kiHJkRRKSiRtIVWQWkhXSIOk92QyWYdsR15O5pM3kyvIx8nXyEPkzxQFignFnbKSIqbsohyhdFEeU95TqVQDqgs1miqi7qI2Ui9Rn1E/ydBkzGR8Zbgym2SqZdpk7sq8kSXK6su6yq6WXSdbLntS9rbsazminIGcuxxbbqNctdwZuX65SXmavIV8oHyKfJF8k/x1+TEFvIKBgqcCVyFP4ZDCJYVhGkLTpbnTOLSttAbaFdoIHUc3pPvSE+mF9F/ovfQJRQXFxYrhitmK1YrnFCVKiJKBkq9SslKx0gmlh0pfFmgscF3AW7BzQcuCuwumlNWUXZR5ygXKrcoPlL+oMFQ8VZJUdqu0qzxVxaiaqC5XzVI9oHpF9bUaXc1BjaNWoHZC7Yk6rG6iHqy+Xv2Q+i31SQ1NDW+NNI1KjUsarzWVNF00EzXLNM9rjmvRtJy0+FplWhe0XjIUGa6MZEYF4zJjQltd20dbrF2n3as9rWOoE6aTq9Oq81SXpMvSjdMt0+3WndDT0gvQy9Fr1nuiT9Rn6Sfo79Pv0Z8yMDSIMNhu0G4wZqhs6Gu4zrDZcNCIauRslG5Ub3TfGGfMMk4y3m98xwQ2sTZJMKk2uW0Km9qY8k33m/YtxC60WyhYWL+wn0lhujIzmc3MITMls6VmuWbtZm8W6S2KXrR7Uc+i7+bW5snmDeYDFgoWfha5Fp0W7yxNLDmW1Zb3rahWXlabrDqs3i42XcxbfGDxI2uadYD1dutu6282tjZCmxabcVs92xjbGtt+Fp0VxCpiXbPD2rnZbbI7a/fZ3sZeZH/C/m8HpkOSQ5PD2BLDJbwlDUuGHXUc2Y51jhInhlOM00EnibO2M9u53vm5i64L1+Wwy6irsWui6zHXN27mbkK3025T7vbuG9y7PBAPb48Cj15PBc8wzyrPZ146XvFezV4T3tbe6727fLA+/j67ffp9NXw5vo2+E362fhv8LvtT/EP8q/yfLzVZKlzaGQAH+AXsCRhcpr9MsKw9EAT6Bu4JfBpkGJQe9Nty3PKg5dXLXwRbBOcE94TQQtaENIV8DHULLQ4dCDMKE4d1h8uGrwxvDJ+K8IgojZBELorcEHkzSjWKH9URjY8Ojz4cPbnCc8XeFSMrrVfmr3y4ynBV9qrrq1VXJ68+t0Z2DXvNyRhsTERMU8xXdiC7nj0Z6xtbEzvBcefs47ziunDLuOM8R14pbzTOMa40bizeMX5P/HiCc0J5wmu+O7+K/zbRJ7E2cSopMOlI0kxyRHJrCiElJuWMQEGQJLicqpmandqXZpqWnyZJt0/fmz4h9BcezoAyVmV0iOjoD+aW2Ei8TTyU6ZRZnfkpKzzrZLZ8tiD71lqTtTvXjq7zWvfzesx6zvruHO2cLTlDG1w31G2ENsZu7N6kuylv08hm781Ht5C2JG35Pdc8tzT3w9aIrZ15Gnmb84a3eW9rzpfJF+b3b3fYXrsDs4O/o3en1c7Knd8LuAU3Cs0Lywu/FnGKbvxk8VPFTzO74nb1FtsUHyjBlQhKHu523n20VL50XenwnoA9bWWMsoKyD3vX7L1evri8dh9pn3ifpGJpRUelXmVJ5deqhKoH1W7VrTXqNTtrpvZz99894HKgpVajtrD2y0H+wUd13nVt9Qb15YdwhzIPvWgIb+j5mfVz42HVw4WHvx0RHJEcDT56udG2sbFJvam4GW4WN48fW3nszi8ev3S0MFvqWpVaC4+D4+LjL3+N+fXhCf8T3SdZJ1tO6Z+qOU07XdAGta1tm2hPaJd0RHX0nfE7093p0Hn6N7PfjpzVPlt9TvFc8XnS+bzzMxfWXZjsSut6fTH+4nD3mu6BS5GX7l9efrn3iv+Va1e9rl7qce25cM3x2tnr9tfP3GDdaL9pc7PtlvWt079b/36616a37bbt7Y47dnc6+5b0nb/rfPfiPY97V+/73r/5YNmDvodhDx/1r+yXPOI+Gnuc/Pjtk8wn0wObB7GDBU/lnpY/U39W/4fxH60SG8m5IY+hW89Dng8Mc4Zf/Znx59eRvBfUF+WjWqONY5ZjZ8e9xu+8XPFy5FXaq+nX+X/J/1XzxujNqb9d/r41ETkx8lb4duZd0XuV90c+LP7QPRk0+exjysfpqYJPKp+OfmZ97vkS8WV0Ousr/mvFN+Nvnd/9vw/OpMzM/OBNzFBbwpj3JR68OLY4WcSQGhb31ORUsZARksbm8BhMhtTE/N98SmwlAO3bAFB+Mq+hCJp7zPm2WUDgnwHP5yFK6LJCpYZ5LbUeANYkqpdk8ONnNffgUMYPc2AG8+J4Qp4AvWo4n5fFF8Sj9xdw+SJ+qoDBFzD+a0z/yuV/wHyf855ZxMsWzfaZmrZWyI9PEDF8BSKeUMCWdsROnv06QmmPGalCEV+cspBhaW5uB0BGnJXlbCmIgnpn7B8zM+8NAMCXAfCteGZmum5m5hs6C2QAgC7xfwAKP9n2U7+jGwAAJfRJREFUeJy9fHm4J0V19jmnqrr7t9999n1hYIAZhh1BEURERYISUUNcYtx346OPSb6YL4aoJCZRY3CLCoksElxABYKADAIzMDAsM8y+L3fu3P23dndVnfP90ffeWZgNJF898zxz7326u87bVXXqvO851Sgi8Mo3YfAECAf+HX4BABzp7wDAXlJEg4AIAECvrGX4ygIWEABAQAAQAA/e+zjhVuJHrK9b32hJzOxEHIAHJIORwlCrfIDlnMoHuhyonAblEQkYgQUUAhzl1byc9ooA9gIgQoRjZjlOq+m+0aS3ZvfFMpxynEIiwIiMAgKCCAAkIoggwgBIokkCTSqkSklPy4ed7cGkgurMkIowIAgg/cHI/0DAIuBFkBAB0Lp0NN03mGyp2u1NaYh4JMasiUJAEJBDDcbxzoVERACYhVkI0IeQL+vJ7eGcDjMnZyqAIJIAGvzDJvkfBFiAs2UWu4G+1q7+ZE0V+hhESaBQAxCKZMsVhcduOQBYGEABZst5/O8EAIhsEQRIpEXMEbS3m9ndufnt4XRCAwAA/LLX9ssDLCKMqAAgsdU9zef2JetTrqFCRA2AKIIgAgyA45Yd0ouAQiCl2FsLSE6BHn8h405OBAgBAYHBOkmML3Tq2dOLp1aiWZkFgIwHnv+/BVgEHIBCIM/VffUNW5I1VgY1BgpJBBAEWITggJsRQYBssQoLIiKC14HqG+385M3xZ68Yfs0Cqra0oGgCADxk2me2kSASMPuEQU1Sp8wsLykEkyfm10sCrF8iWkFQADISb9k2uqoqu53RCgPwmsUbBovCOaOsoPeCgCyglVWgEk+5QAfapc7FaWAoioKgd5Bavg7kKwUE5EZMIg4FBTSgR0AgRiTvAQVAIeWBuNc/Nzi8Y0p+8ZzimYiRgBNQJ+7MThSwCAsAATlIdo2u2hs/7dErFSjO/I5nBEcQmMhs2jfa06YK2rOjvMnvq2mt7PS26P61+R//3l99VvVt51VuWk73Pm+G41Y5VAztt6xE72t/dCaDz+VDr0zSapJS3osASV5FSK0kMd5rD0A5D63trSebzf45becWwynMFkif4NZ1goCFQRSqON23YWTFIGxRRmmvRURQAEAABMEwVr5wm356C97xuaqwhKb0wr7o2/fntg4OffNPtGDeir/xf8w5C+Ubv8EffUI+8h0VquJ3fhs8tF5t3+9LAewdiv57Vf2d5+OfX4QDVd9RtgDFxzdLI4kuPMmikAAwI5qQ/ABsGhkdnFu8cFp+4Ylv1MdZAAwg4ABYgRqKNz87dN+IbAvJKG8cCgKgAErmxSQJlbtqWbRnSPUNA5HJRYW7nikknozKbRvULedDZSIjQ1XVTIJ603mnDOWWvxB2FIavOSucNzVct0eVC4W2YtuXfxFc/5vgrtX5Slv5Bw+V/vPxtBQ6YO0l1gDILKghcqq5qXrv5pFVAA5EBI7vj44DGAHAE4DqjV9YM/pAPRgiZViQiT2wJRCEWIkvhhGoxKbNN5yanjG38L0HdT6gxPK8LvXAc/5v3yZLZ5kfP5hbvzfeNciNlM+cj5+9CWN2/Q3srLi+mr76/Li9GDy8sfWP72pddor60e+62tv91K7C6u20fKN85k0gTJ5bBWpDDcIIAgKBJ614d/zw+qHfOvEIfFzAx/bSAgKAuKe5ZlPtEdCJEg0iICCGgihsNVuKRQUmt3J7Mr+7VQ6DQli5b13lo/8xsvYr6T2rK2TgZ0/EWwbSD7xGL5vLqzbrk2eOzOoomMhWq5ILHUopEfiPh2ThZM6H9O171Nff11qzbfLu0eZ37/ObvxO+71vhqE+TVvVTl7TeemHuFyvUbY80v/PB2KYKEEFACEgSn/ToU07uuFRjKCB49Bl+DMBeBBDV3tbaDbWHFTnMdhsWMVTorePNv5OPXTHanW+/9XG4f62844L09ac3aiMVpbsu+arrLrcmlVrvfTWfPTvYOuimdzZypALDznvrtGPQClgAQBColKPUS2xNPUEv6scP53pr2JEL501Nbryv+KNPurdcX7/x/fGkcm7VtubUcvyqBSblZoCGswVlkNj6epc+fVHHxRoiwKNiPjJgAQZhQL0/3r5h9NeOmgEUBDwAoBeMjF63v/LBH+bOm7v3+x8qfOZHhSe32uG6fPyK/e8+T5wt/vSJwOjmFWf4UFubQmgwdeQFRIQQM3NExliGCHhBBEEkImcoYtR7h92kSuGD/xGdPgdGGrx3oHnHF4LX/w1Obh+8+cO6Go9wkk9FCGncfERIfDwpPGtx5VUIGvDIq/WogBGoavc/P3RXqmINgiKZeyIBZq/burp/8VTunf86esPb4/e+2vzlreVyqb65t3HzR2wz8eVAaSW1BL14hcRZ7AECQAAk4OhAFD1uBwCCQwidJBqU0l7YCBR3jujP30p/8aZgz4j7p1+0Hrse7n+GfvVM/M13N9Bj6vU4RQNARLI2nVE8Z17pfBQWpBcH3kd8DYxCqa9uGHk4xdhI9sgxloYsLYM0PNpc/lR8x6fkBw9U7l8P1108tGsgvfY8BmEFVE9luMUgXsFY/DXO8QTA4+FoYfx1MEErQOMRrFeOveeRuR2N2z8RO5d88Rb31Xex58qX78rvrJY/cVN5T1WTFpmIRgEYA4N7mqv6mhsAlBf3YmxHACzgAZOt1RVV2BEACYggYIbYM+aCnrDo2o1ev1taafFr74H3/DvsGPJfeXvzitNjawlRCL1GElB8+KOZWeSovjTwYDwwZRAQETFxsbfpBQviOz/jFkzrfMvX7ZvPwFs+Ed6+wmzp9zkDkpGTbI54EUUEW2sr6m6/QgMv2qgOn9JerELV39i4tn6fUgqFPSE5YQLxEnRVCjc9lv/hg8n33i+DtcK1307bA/vxN9auOTsJkFKvAbMAVw4dRkYgAa8gHwXekk+SFD0CEIwFhXTQxRNuVsaf44SLFro+/l/m/udrN74Hb1tJU9qH//Halku1ZKECiWYlAqIAyUvSQfNO73wjCQAigDoaYAGR2NefHr7DQY1AMwp5UYVczJYAKr/bHH7jnmjRLL57RXzjn1G+2OofiC8/NfVsUgYaQ/uiKQOK0RrGJGr98pmRpzf1fPa1ZlZ36FUjroOAQkJhyQQBBEVGvJu4VwBCUrmP35zfMdC86cPuyc2qo9S4YD43Y6fQRzrQWiNCK029z+Y2kpcU5xcumlZaBuIRD5AqOswyQNrZeLwhQwiBgKAXnzPR3atLq3aF5Q7/1Lbcmt3y2sX89ffgJ/+zVjG1a852sVWehV68MA80ixwaVX94a+PWx+PfrytO7nrsl09+60vfzVGktbbiSOnQ5MMgYgr2D+9jZhiTRUCxJBrSv7169G+uGp1Wqr/1zMHXzGfvVHsxHwSV3bXSvWu7P3ZT5fq7TRSJCAIgK9S8s/5M1Q+gKDkoIJmIpTMmRNV4b2+yNcIIxQMAg5hA83M7wnJONy1/4ara5aeaP/k2vu705Nefh44S9lVJE8DhUsbBDQHBYDqc8opN+f0DjfMXeIge+MXyZRctfmHf0+v61o/ycFHlOqJurc3uWm8F295y2h8lPkHMwlekxLV6ijilQqMxCmDeBM/tNis2BRv3iYUCheFNjzZ/9VknrMelAUITq+q+2vOVtotBDrApOjC2DKn4Hc3VCFUZn/QayI806l+5Bl53evCZH5fe8jWZ0+O+9s50Srkxq62GDrSSF3uGQxt7HwS5ePVAfuv+3oGhKW88b+Xjq9RUqLwW7tp255b0+SHZsyPZ9vTIEyuGHt5Re/qkaYtZHfoIRLQMrRQUiUIE1PlIbe4LP3qFyZvggadrX3tH7Q2LqRED4cQAKoL+1oZa2otIEzZOAHYINJzuGrC7CUPM6EAhDB/b2PW52zo3Dw6dOtM99Ld42vTw7P/T2tPf/OurpNoc06mOCBKBETyAF28MpkPNxlM7pX+g2l3sueikDc+vm724Y9WuR4BSUASilGJUKSVywZQ3L2qfb22MeOicQQBCQQBCSG1tYXfzhnfWbVp+fHP9tGkjn7zMjjYEMdGgOFMGRIlOdHN3cy0cNCRjUxpFCclAay1gE8AAeAEB7w2pyk9X0cObgjcvST506dC/vC+cXNHnnjScporwaGizDlg4oMCQToDr+4eKubB3294pH7rStweDgyPvvfbapq6/sGf1/nSvTRMvvtN0nzz91PmTFsfeHiMYHgvQEsuFvLrtsda+0daDfxmMNoOBemt2Z9RMEyWGx5YYEQ2m6xvujILuEmghhBlgQaSq3TeS7tJEY+seyXpOrzhj+Ecf4f9aXt49HFz+NbN4RvOb143M6wpGW0BHUCEPmCUS6qjWqj2348kttc2ekkvgtL7Tz4vOk3LvdmyhIvSpragOo8OOtrb5k08p6zbPScyWAPGYy4QRFKBP0+anLoPXL839dg1c/3Oa0U3nzo0/dKlTYxEOo1einaS9rY1zSx0oBhD1+DuDodb6RFIDCkE8CAa6smsY7KB99YJgtJEqrT/y+uTnj/KUdqy18Cg70JhBIrkgt6N/670v3FWXAVUOptTz9LNHR08tbC3179iwrX0ZPLzhV/vsjkRSx2j3+bW9ay4/5c2dYQdLJmXKMSh9FrZh6uPuskpdx7nfaPzzn9YW9ZjP/5TfdYHvKijrASBjO4RqJN5mC0sDyguIBgAE9JIOxruQBASEhSJd2Njb9oYbyk1O5nag0fGze5tfuiq+4T0t19KIx+CUIhLoYMfwzp+tuQnzUjIdJsovWOnc4IC5ZEZ7PucoaVtMraRRiXqct4lvxmy3JS/c+fTANWf/eSnMi4XjjXGGG8V7nUj17IXVq85RDOm/vFsiMATIImPeSxCp4Ueq6e7uaKEAEIMHgIYdbPCwRswWCAnGVqpffsfAH51ZP2mGXHeJ/+lHZVKlZWMNKJkocBQrBEEQVmxdzqErB8WI1CRb6Fle719UpqWFkKlsygXOlYNSGAShzocqF6HujLr7eeDB9b9CBXKEWPvIA60TNzS3s/lP17X/3Z3mH+5UI7WOz9xS3DnkQ40s41ehp3Q42Q4CIKJFBBCGkz2MbEQzCCC6xOrTZ8UXnGzfe6F/bFPzkU3uoV3woUtVpH1sCfFoFomg0mbfaO/+xvZ8qSMARcV891pJNgzJp08L83nMskYSeHDkYwceMScg4uJKrrxjcOveob3Ty3OcawEefU6PNyUEzmHJ4PJ1pRk9+q9/JnuH4sEmzewGcZniKwCEWE37LLeMymkAzWKH0j6a2KKYdSGX/9mq6JntLjLub67RFy6W79yblEKdeHtMOxh8RLk9td2MNsIymrCsS5Mf6RvozpuLpgdOVKCEQAScOERy4kBQWISFmRvY2DO4fXb7HHcEnnOE5gmClk+nVtyVZ8BtK+0bTk+vva6xdBbUGmOOCwFACKku1abrq6jZWgG0vG1yP5IwAKMAInjWk8rhE1sKa/e6321wlQKfNQMntbUasTnq6AIAIDAD1xsjSutAtDb5zn3Grtpn33RKoSsPDU+RJgBmTlkBSiKpA6/FKTaKLRkabg0Jixxdsji8R0Jrbe2Np7VfepoNMFi5zeTCdH4ne88i2mWiCrG0WmmtEoIGhNiPWG4qQgDRDB4xtqm/6JTaQ6eY7ftLu0bDW1cm8ybFIEw4lj85KmAQEPGiSBMpE6meFYPNGMI3zGbAQlhUhlCEGcilIsKKHSunlCIipbQynt0JaI8HmhBCM3VLZvCnb83f82zrjUtLq7a4v7469qnPGV8poBsLpGt2cHIWeLR8P0OqUYtIolAxFwTD3sFwY2+wZ5i37reDw61zLwX2oT8RARi1ViBAhioNpR7uj8+bHs0JTAImivLaCIBzDgQ8W4uJAiRAIiJERDHKvKSUKAo4TaZpZd1u9ek3RxefbD76Q377N6N6s/FXVzXedYGMNhAIEOvSP7Yt1bmO6EA0CmhPaLD4uZt7bn/aT++wcYqd5dFZHaiDIBF/xGz+Id0TCraX2nEIfajKzzTdzhp9eAkIKiAFolSAgOwZkQkFCDOahTKWOKtEnYhjis+JNEHQll1ngV6/RP7+Trh7RrJ4qsybki6eJoumSDMGRSCggVJIYm5pALCullFuySQ1JvuBNw6rSGnNcZp88ap0UputNRyw0CFM/Ujdk7hkanlmSCXj1KT7h/rnFeT0XBLHWuet+MTHiGjFOmEHIiIgAjhG+I0Pu0vTvLzEVCiSj7lx3gK+dZ7pLjcKqrVgeppYilO0buLFee+cr2oWZyWZwEHInmV08ZTGN9+tn90R3vZ4+KX/hnNmw/svkVYL6HimIHh2bfnOaeVpdvse9Uy1+WczMUyiJlq0GlNkQQArPmGXsrPeexFmFu9Tl+aj8qTyZOddFreeYBMSTB3Pn6w//H3zzK7KjC5894Vw3atS6wgBxxRCxRynvkZOvPPphEsUAAVCrVjqCS2dSTdcl7ztbNNW8oQkeOzhBQAE8QSGcdb0ZZMejUdznJzbHjdaFpz1SWxbzbTZsq3YxjHHCSdWrOU09S5B1bLpzM65hSjP7I5Frw9DC2AccyUK719jHLu/fwfOmlT8zbPgU4UH5iMKCqROfBZLpwdn1UQEAp3bNlD67E/iK05zV51Fz+1B5hN840qgSf5kNaX36fSZc/L1zjTXshQoBm/Rk9KIwMKO04STxKeJT6y3CTcLvnzGpHOdOIRj8bAXvWKwmlQjbS2dkQYX0PkLo9sfk/dfbHOG6nY8wBy7ltlnpRUMB0euXrAY4V2ri4VSZH3PbU+qnzxCoTmR1DkjOBDMB/LUlqH9tcaF7a5eS53ENo5t3HJxkjbjpBkn9ThtxGkS29hZm1ArrVfPnnZWV75DrEM8forosObZcU8pV/c00PDf+0Bz8WTVcBMTUhAAwSN7tHr8NR3UFGIt9m87p7Gh1/zwEZc6//U/9ak9nJEDgOC45nqg7sqJlL0M3r6y8KpTTjpv6fKNv+qstBNH1rBGa4QAQcB7sCmLTW0so9Vq9dT2c06efkHLVwknSiROdFYrBq/IjCbJHQ933P4o99f1tEp8+6edpHqcGI+J1yKasgkkcDA9YQAUoctOr33patddNsMxt1JUCAeRNkZB7U0QKcx0ZBLwIIw5E22vDmzrnfuv78stWFCtjW6oP5FQGjVMokMxQIDkvAObgLMqgQSXzbz8ggXnS5KIB7EeWQmRoBACATISCQsIHyX48ggGlR+ulhfO1pcs8Rv2wBMbNYLnQ1wBAhKg1kKEhpFBxnQkEVFhED6/S//lraVPXpE8t8MCm+++P05SNU4bBERQUR37X9iaxJ7JMIhSgOBUITe6arft7hqsWXvPU3Noat+urmpPLZksCmKwyKmoSBLtKEWz089IZi3iSfv3bYp6isWedlWMJERBFJc6m1rH4FKLhET6KJ4MWTwKMOPOgfrPH49e6OVLT/aRgvqYM8hsRiFNkUYyhIVDljagtzZdODmZ3aG29um5k9O20McWD9oqBEQh2ard+vTu3Y9uGdywf2C0Vqs1bZJa8R2EOgxH7n4iQAhJown03Fy0pBQqiLc1pWbNtLwqKaha2VxvjAytgUdFS74YtXV3tc3s7JrXXZk7tTJrUnFaV6GzSGV2hGkzZXdkVso5Q4qi0ZgXTS0sm8/BWgAQrQ7NjwuiUlRAEb9h+Ld73RoNAYL4MTUBXaSLw6nS2uwcoL3DI6+ea1kRsGIZXxWiSOkwkqYfWL19YGU63F/d9cK6hP3+HduVs+QFibNlICn4xDFBoFABMjMBEwEGIloxsfKaPINzzJ5FlPcYatVWKHcUuk6a1jV/xoy3nBFN7xDrJjAjgENQIKWN/fWTewr/vSq6e7VGkP5G85OXj1y5RNeTCRnAEhjJn1n5Yw1AoSpKCqAABEgg0VCsp23LNwxfML/9t8+3ffRHdmpF/uGPB96wDKuNCekeAZ33xF63Uffr506bPnPutAu3bVyfJHbvhu1GatMXzLVpLMDCAoiSJZaARRgYgIFFvDhhRkFhzoTxzE2IeHE+SVPDtGvD2o0///WUcxaHc1SSpAeCnyx3Y4H/7x0d179n5AOX2Q9coQDZtlyjha3UKxxLiSEw+EDyocppADCqMqbGCbBwkC+qnz9KD60xFy6snzzF/9v7zGkz5MZ78pctayLSgWktWustD6zljY3d+4eWnnXVgkX56ScvGxiqYWGWH9mx9OJzPQBlGr8AM7CAH/eZLCBeCFEYGMQhM3tgLUwAAuhFWBwX8lHxkceeWvnYxge3nbGwXefVRDzACMZJGmnz3tfKd++ddEveWN+sjlqW5v95m+8pmHTMxwoQic1jUZHRAJA3FYVBVjaW+e8QlTtzLnRPtrk8LVuU//ov7VDdaQBRgA6y8jEBEqhUutbt3bf7md49T31//YrVS869tG3xkjAsjdqUxXow6bjbEAFhHhPWWMgLCwPoLO9MKMiQ6QAAICggwJ6biTdhMJxKMclVd/X2nD4vaY1J1ggAiJy4+MozcHZ3/jM/8b3D5SvPGhUL7UWdMsFYioUERSAKSwCoASBHlQiKKVQRCIhss8WvWkQf+0FHPcXp5fCe5/FXz+MPP5yIC8ROFFAgoneu69zJF1444+wNA5sf2LT592vu+vGz7T3TO+edPG3G5I3P5qwX9CAgJAwCzMJEkEUwkv2fJdmF2IuIB8xySgyYqSI6IOWD111yWcxbijMX2NQdFg4QYKvZglOm2t9+Qd/5JKzYbK8+kwONrYkFDIIePZVVNwCgF4cia4fu6+P1ARqHQh5VFFRW7wq/+sugr+ZmdrsPXNp38Tyst/BQ8pBljD2yilQQhPGeZOPyFxq9Kez3u1avMZFh8SCA2ZYPAiIKSEAIJBM+BQURJuQrFEIhABH0gECiGGjRa169+YnHZ7/5pEXXXRhX60RH2Jy0k6aRYqmkfrpSTeuonT2Lmuk4YBR0wGpZ5e3FcBI6sQr0rtpTG+MHFZaUWARgES6YSBlqOckFLbZUT+jo4g4LeOEgNEGoXD2NmwwtR2NSKwMIUDaiWZwlgASgssIqREKCsQTfgeiVsnIDJh83moEKOAqiUv5oUWeWEBdm15YLmy5lN5E+Q0AHtoQ9SzrersloBAUA7cH0oFlwmpFRUBCR6i4Bi4S+mmQxCosgi0CmH6AAgmQCHGhATWStTxMHmqKidm3hWFmwCAuTp6zKhBEEEUlwIgoay52giBw0XcdiVQIq9nSKOGNd4p1iBIXER2BTgoCKdDW2BErw4AtEpMNM12QERGdvomg6i1QelAGDKuvNayBBBhBFykmCGAaG85pABc2WTrxHJaFykbJKSSIcN5G9IRBmL4guc/uYhjoohCCSeSLUWiWW6wlnjufAnDkMgRz4AYFYXKCDthLHCbas1Wj8IZiznz0CjXV7YOQBRbNuD+dlfxlnhaQ7goUj8X5RBoQRQDMIgBCEKTfzOq/D3Jpd5Xufo6e30EjLkXiNSlCigl4wNTln9sCr5qSFAnpGYMCxOmLImdzGHaX/fFTloxA0a0wbLdfR1vzga5kED0mnHG25ILBALihuHaj82wOjV55au+xU1UgYj7CUM3Z0WNTI4ko4qRh0CngEpcc7w878/N3xsw6bJGpsNBDQiSuEnbtH2/7+lsLPViRKuT+9qPbOU7inRAQ01Aif2RXevDz81m/qD38xXtpG1VjUONEhEC+mVOFLl3Z+8fb2dX0OoPcjF6WvPjnNykSP2cZKpAgYRIVR+Vv3t//okdwj6+C38xt5A8wnwqY8ofO2I5yvVY7B4fgII4PkTLkSzO+3TylUWTUJiHCoy9uH2q7798nP7Rme1zly+6dqy2alSayteAQFyFef1X7l6W3XfFP1dARxgjiWBRMEAGTv6935YNG5nb98Wm3sF4D0mgtGXnOS6huEbO0c02hLACIqH7Y9tTN/62NtWjc37Y+//3D8V2/m4TooOvbtguCQA2jrKi4EAAIFEwlxBFagp+UXaolkfGExQB5V7u/u7Hlujw9140tvry2Z5vYOqNFYWoluJFhvuYHRWk+576vXpAXtRfiwiYaAjrHZsM6DZ2CGWhMa8fGHBgAAAgblxeZN8Z/u8mfOG5jbZRDDHz5Q2NLHoYHjCRIISrg1yZyS1yUZJ9gTgElE2oIpXcHMGCyAIWYMjNre3/Hg+pCwMa0y9JqT4loDAo2EgCSAHlF5dpGO33RWk8x4KudAoyxkpQMVB0JIkG2+x7R1fJuJ26KehzbB/6yufud9o9eejSLte2u5f79f5UMWOUYSBAE8sMH2qaW5cJDScxD3QECkqfnzI68FLAihItVfC+upY4H2ggoD8pCRUhSWYgBdRWjPSzmXGogqeUWKX+R7Jurkxnnp8aAehNkSlESXv3a3fev5tVPnj77v4r75XYTYdfNj+We2Qz48WhZkjP/6ZLJZWDBTBA4kEA5odwgoIJWwc3qwdKt9kjASAckZb0hboNEmxpbL5IRRgEJtfrdWrd+XK5VUoIR90mz6y5ZAT1mszzolwGOsMcx0Fc58I9K4VuQIVHZ8Q1gV8vnfrVfL15kffDC/Yo0vKLhokd/8+1w9rfzzb5o//jizZQ2KwRMQjyncjKJFJWhz0DmrsEQgK7qEwwHD+DhMKy/dM7LVyjA4tnN6GnM6Smv26V1Dau0O9drFwWA1jsgR5r2KYszf8XD7yi0I0P/WM+2V58UsykuaIyTlE0tj9bKQhRSCWaHuxPYrXDTiQVImYBdoDaBiJ6RIWACUQPlf7lXdbfEtj1SaMSNaY5K2XL4ed/5sdeNP1g9dtkCPNmxAxGI1gdGmZQUgVYiOZxWXhbpNgA8+AnKYhIAIGKjySbnzjVUiNq0UWh97Yy9BMeEp19+dayRJZ944kWZSu2je8F9dbd92dqhV3uj47efG0yrOpb6kC0Oj7Vv6VC5EUQAg7LOqEUAETaxRNMYasFJou+OJ/KMbVKQ4Crp6R/K7+sKQPHGTvKvk2+5dYx7d1H/HJ0Z+8on9t3928Cef6rvrL/o/96bYSY6h4+t3BanzGdVQupCm7Rt2a/KkgpTTybRwcu5UyDj2Qe3FmomI+O78/O7odHbgG6197z7Pffna4XwQPbZp+tX/2vN0L5XypUpZ8mHIqRmoWueb1mHfkPeewlx55faZr/v6jIu/VrrxAVcOwTGEQRA7EEEWXXflWpqvpVEj6dpdLX7rvuK+qhSLhd9v6r78hmkX31D68SMqDCIVdDa48KU78HWL0wsWjHJaN1ALKUlc7aOXJCdNUiKVR9Z33fK46akgIDXSSe+6cfprbuj6yE3Atiyl2ZUL+EiC0ItPtVAmYs4pn1Mf7q/xTmjKvk9fGl04r/N7DwX3P1u54vrikpnpnM6KIjVQ0+v3NxZPry6ZXj9rNjVTDg0NNWn7MAgUtg/V2EsuDO573q3eGXeXWhpLX/l59I27KXEsHI7EttocmT+FnC/01czO0QBgdOdIkMuZe1aZG+/P76+NVCL16ZuKn7g8qYREZPYOB9+812lV7y4BoP7nX+c27pXPXxk1HWwdzDft8Pb9YROnT74oCjq8WBg7unfQHD5ygbh4FNXwA88M/yKhGoimXChG2naP0tZ+vXPQ1BIfamwrJPMnt2a2+UpBYp+4lEAoNB2P7eDhevLaBSOhGNFR72iUMEaGUTCxXpgYwYP24o1uTS66EMnotse3czUZvWQBK6T9tdxQA7oKuubS2Nk5HWwQALxNws1DulzAiARAkoRqrr6gRwemuL0aPLutefbM7jlvmhUtcchKAF9UF3+0Mw9jpZcjyfbnRu9hSoxTIuwijYFBo1mB8iIsGFuwTjwTEmTyjwjnQ9BEtXiM8hgliCJMgoLIBONsC0AEUi8oKGCLgRHCegwgYDRrEueBCBSpls2iXSCwOaNT9iCIyAqMEFuHnn3OWE2z3MKFhQuFDB5Fyj/WqZax+q3m5ufrvwFyxocefXY2K6t/zjZvAESEg8u2kEWyWseJB415aIBxKzJi6zPwAIwALNqLEMKE+D1x90G8kVjGazcgI5kI6E0aWzeNli3qfBVBHo9eL3Gcw5YiHlH1Nbdurj7gVFNL4BEBPI5v5QcOAL+UtP3BN2b3yvimSHLI0/CgX7NrVFYXekAh115765NptHRBx0UaAgAGPCrk454u9QKAoIbiTeuqy1McMhhlpYlwAqei/rebIDKi9X56uPSk0nmEwXFrB07kOO0Y5lqye93I8ib2KtLAEWBy4vmuV7oJZGdYkHUKM/LnzyovAyYhoeNH6ccHLDBWy4exizeO/m7IP0dUACFBT/8rH4k4jjEE4lF7TgpcmlN+dXd+IYiMH5x5BQBPdCUIwOL21p7bGq9IVctIRPL/eW4zgnKALOkkmjuncn7O9DCIOuG59tJOiGdaMiFWk97t9ZUDbjcRaxz7egMTA4zVsL7SbWz4BD07jrA4O7dscuEMQBJwgOoEq9jgpQIGgKwkAhFTiPc2N/TV18XcB9oLBigKwR6LIr2shoCORICBUyXB5OC0mcUzI10ByBJi6viPOPhpL+8rDwKORQGi89W9zS1707VN7jfgFIZZQmx8I30JifzDLct4NAqLs+JC39YdzJ1cPKUSdoGQYMYKXvLDX/ZnLXjsDJUQIKTcGGru2NdaU/d9DiwrIQo0EwJ6EAJm0AKA6I6iUWTZM0NCBCwkImjRsjjtTQ7busJ5Pfn5xaAHQLEwgBC+tIH9wwEfsFXGfaMDW033DrZ2jqa7Yh5IUABEAwIaIAtAKIoOO4eXMWQER4LCkn23hYmQc1Ap6ykd4az2cHqoSzAWBb2cUT2kw1fkWzwCToQRdBass0/qdnTU7a7afS1XbWIdOBbvBHk8MDhYgEdBZNERB4qCULe3qcn5cFLZdEeqCNlR8LGkxCvgHV7hjw8B8NhOOd6EnfWjDW4kvma5lrLL0uJZvKgAFeUClQ+oGKpyqApGhQcBY8kKvF+5COcVBzzRZJwqHMXWo7izMXvw+CHEy2v/DyDV9XpRcYScAAAAAElFTkSuQmCC"

SEC1=["Adopts objectives of syllabi and curricula","Selects content and prepares appropriate instructional materials & visual aids","Adopts appropriate teaching methods","Relates new lesson with students\u2019 previous knowledge and skills","Provides appropriate motivation","Presents and develops lessons","Conveys ideas clearly","Utilizes the art of questioning to develop higher level of thinking","Ensures students\u2019 participation","Addresses individual differences","Shows mastery of the subject matter","Diagnoses learner\u2019s needs","Evaluates learning outcomes","Differentiates instruction to suit students\u2019 needs","Uses a variety of effective instructional strategies and resources","Involves students in cooperative learning to enhance higher-order thinking skills","Uses a variety of formal and informal assessment strategies to guide instruction","Gives constructive and frequent feedback to students on their learning","Provides independent practice activities","Checks for understanding and adjusts instruction as needed","English oral communication skills","English written communication skills"]
SEC2=["Maintains a clean and orderly classroom","Maintains a classroom conducive to learning","Establishes clear expectations for classroom rules and procedures early in the school year and enforces them consistently and fairly","Maximizes instructional time.","Establishes a climate of trust and teamwork by being fair, caring, respectful, and enthusiastic","Respects students\u2019 diversity, including language, culture, race, gender and special needs","Cares about students as individuals and makes them feel valued","Promotes self-discipline","Manages disruptive behaviour, distractions and misconduct well.","Keeps criticisms fair, objective and minimal"]
SEC3=["Supports and participates in parent-teacher activities","Works well with the other teachers and administration","Has positive relationships with students individually and in groups","Provides a climate which opens up communication between the teacher and the parent","Uses discretion in handling confidential information and difficult situations","Informs administrators and/or appropriate personnel of school-related matters in a timely manner","Cooperates with parents in the best interest of students","Acknowledges rights of others to hold different views","Handles information about students\u2019 private lives in a proper manner","Has a positive disposition and attitude"]
SEC4=["Adheres to school policies and guidelines at all times","Exerts effort to improve one\u2019s knowledge and skills","Takes responsibility for his/her actions","Channels concerns to proper authorities and not in social media nor to other people who are not of concern","Observes professional ethics","Is punctual in reporting to school and in submitting requirements","Is not often absent; when absent, only for valid reasons","Willingly accepts and satisfactorily performs tasks or assignments assigned to him/her, like committee work for co-curricular activities","Does not participate in rumor mongering","Wears the prescribed uniform at all times.",None,"Honesty/Integrity","Initiative/Resourcefulness","Courtesy","Respect for Authority","Teamwork","Leadership","Stress Tolerance","Fairness/Justice","Good Grooming","Human Relations","Self Confidence"]

def set_no_border(cell):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    tcBorders=OxmlElement("w:tcBorders")
    for side in ["top","left","bottom","right"]:
        el=OxmlElement(f"w:{side}"); el.set(qn("w:val"),"none"); tcBorders.append(el)
    for old in tcPr.findall(qn("w:tcBorders")): tcPr.remove(old)
    tcPr.append(tcBorders)

def fp(para, align=WD_ALIGN_PARAGRAPH.LEFT):
    para.alignment=align
    para.paragraph_format.space_before=Pt(0)
    para.paragraph_format.space_after=Pt(0)
    para.paragraph_format.line_spacing=1.0

def ar(para, text, bold=None, italic=False, underline=False, size=12, color=None):
    run=para.add_run(text)
    run.font.name="Calibri"; run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if italic: run.italic=True
    if underline: run.underline=True
    if color: run.font.color.rgb=RGBColor.from_string(color)
    return run

def make_doc(data):
    ratings_data=data.get("ratings",{})
    info=data.get("info",{})
    
    doc=Document()
    sec=doc.sections[0]
    sec.page_width=Cm(21.59); sec.page_height=Cm(33.02)
    sec.top_margin=Cm(1.27); sec.bottom_margin=Cm(1.27)
    sec.left_margin=Cm(1.27); sec.right_margin=Cm(1.27)
    
    ns=doc.styles["Normal"]
    ns.font.name="Calibri"; ns.font.size=Pt(12)
    ns.paragraph_format.space_before=Pt(0)
    ns.paragraph_format.space_after=Pt(0)
    ns.paragraph_format.line_spacing=1.0
    
    contentW=21.59-1.27*2
    cols_cm=[12.723,1.263,1.422,1.106,1.265,0]
    cols_cm[-1]=contentW-sum(cols_cm[:-1])

    # Header
    ht=doc.add_table(rows=1,cols=3); ht.style="Table Grid"
    for cell in ht.rows[0].cells:
        set_no_border(cell)
        for p in cell.paragraphs: fp(p)
    ht.rows[0].cells[0].width=Cm(1.8)
    ht.rows[0].cells[1].width=Cm(0.3)
    ht.rows[0].cells[2].width=Cm(contentW-2.1)
    
    lp=ht.rows[0].cells[0].paragraphs[0]; fp(lp,WD_ALIGN_PARAGRAPH.CENTER)
    ht.rows[0].cells[0].vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    lr=lp.add_run(); lr.add_picture(io.BytesIO(base64.b64decode(LOGO_SMALL)),width=Cm(1.8),height=Cm(1.8))
    
    np2=ht.rows[0].cells[2].paragraphs[0]; fp(np2,WD_ALIGN_PARAGRAPH.LEFT)
    ht.rows[0].cells[2].vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    ar(np2,"GreAtMinds Leadership Academy",bold=True,size=26,color="C0392B")

    tp=doc.add_paragraph(); fp(tp,WD_ALIGN_PARAGRAPH.CENTER); tp.paragraph_format.space_before=Pt(6)
    ar(tp,"TEACHER\u2019S PERFORMANCE RATING SHEET",bold=True)

    p=doc.add_paragraph(); fp(p); p.paragraph_format.space_before=Pt(6)
    ar(p,"Name of Teacher: ",bold=True); ar(p,info.get("name",""))
    ar(p," \t"); ar(p,"Advisory Class: ",bold=True); ar(p,info.get("cls",""))

    p=doc.add_paragraph(); fp(p)
    ar(p,"Subjects & Grade Levels Handled: ",bold=True); ar(p,info.get("subj",""))
    ar(p," \t"); ar(p,"Date Observed: ",bold=True); ar(p,info.get("dateobs",""))

    p=doc.add_paragraph(); fp(p); ar(p,"Legend:",bold=True)
    p=doc.add_paragraph(); fp(p)
    ar(p,"5  \u2013 Outstanding\t4 \u2013 Very Satisfactory\t       3 \u2013 Satisfactory\t   2 \u2013 Needs Improvement  \t1 - Poor")

    def add_sec(title,items,sid):
        total_r=1+sum(1 for x in items if x is not None)+(1 if None in items else 0)
        tbl=doc.add_table(rows=total_r,cols=6); tbl.style="Table Grid"
        for row in tbl.rows:
            for ci,cell in enumerate(row.cells):
                cell.width=Cm(cols_cm[ci])
                for para in cell.paragraphs: fp(para)
        
        hrow=tbl.rows[0]
        for ci,(txt,al) in enumerate([(title,WD_ALIGN_PARAGRAPH.LEFT),("5 (O)",WD_ALIGN_PARAGRAPH.CENTER),("4 (VS)",WD_ALIGN_PARAGRAPH.CENTER),("3 (S)",WD_ALIGN_PARAGRAPH.CENTER),("2 (NI)",WD_ALIGN_PARAGRAPH.CENTER),("1 (P)",WD_ALIGN_PARAGRAPH.CENTER)]):
            p2=hrow.cells[ci].paragraphs[0]; fp(p2,al); ar(p2,txt,bold=True)

        n=0; ri2=0; ridx=1
        for item in items:
            if item is None:
                row=tbl.rows[ridx]
                merged=row.cells[0]
                for c in range(1,6): merged=merged.merge(row.cells[c])
                p2=merged.paragraphs[0]; fp(p2); ar(p2,"Demonstrates the following values:")
                ridx+=1
            else:
                n+=1; row=tbl.rows[ridx]
                p2=row.cells[0].paragraphs[0]; fp(p2); ar(p2,f"{n}. {item}")
                sel=ratings_data.get(sid,{}).get(str(ri2),0)
                for ci,v in enumerate([5,4,3,2,1]):
                    cp=row.cells[ci+1].paragraphs[0]; fp(cp,WD_ALIGN_PARAGRAPH.CENTER)
                    ar(cp,"\u2713" if sel==v else "")
                ridx+=1; ri2+=1

    def add_score(r,pct,sid):
        p2=doc.add_paragraph(); fp(p2)
        ar(p2,"Score: ",bold=True); ar(p2,str(r.get("score","")),underline=True)
        ar(p2,"\t\t")
        ar(p2,"Average: ",bold=True); ar(p2,str(r.get("avg","")),underline=True)
        ar(p2,"\t  ")
        ar(p2,"Description: ",bold=True); ar(p2,r.get("desc",""),underline=True)
        ar(p2,"\t")
        ar(p2,f"Sub-Rating (AVG x {pct}): ",bold=True); ar(p2,str(r.get("sub","")))
        rem=info.get("remarks",{}).get(sid,"")
        if rem:
            rp=doc.add_paragraph(); fp(rp)
            ar(rp,"Remarks: ",bold=True); ar(rp,rem)

    add_sec("I.  INSTRUCTIONAL COMPETENCE (40%)",SEC1,"s1")
    add_score(info.get("r1",{}),"40%","s1")
    add_sec("II.  LEARNING ENVIRONMENT (20%)",SEC2,"s2")
    add_score(info.get("r2",{}),"20%","s2")

    doc.add_page_break()

    add_sec("III.  INTERPERSONAL RELATIONS (20%)",SEC3,"s3")
    add_score(info.get("r3",{}),"20%","s3")
    add_sec("IV.  PROFESSIONALISM & PERSONAL CHARACTERISTICS (20%)",SEC4,"s4")
    add_score(info.get("r4",{}),"20%","s4")

    p2=doc.add_paragraph(); fp(p2); ar(p2,"Legend: ",bold=True)

    lt=doc.add_table(rows=2,cols=3); lt.style="Table Grid"
    lcols=[6.509,6.191,contentW-6.509-6.191]
    for row in lt.rows:
        for ci,cell in enumerate(row.cells):
            cell.width=Cm(lcols[ci]); set_no_border(cell)
            for para in cell.paragraphs: fp(para)
    ldata=[[("1.0 to 1.59","    Poor"),("2.6 to 3.59","     Satisfactory"),("4.6 to 5.0","       Outstanding")],[("1.6 to 2.59","    Needs Improvement"),("3.6 to 4.59","     Very Satisfactory"),("","")]]
    for ri3,row in enumerate(lt.rows):
        for ci3,cell in enumerate(row.cells):
            p3=cell.paragraphs[0]; fp(p3); b,n=ldata[ri3][ci3]; ar(p3,b,bold=True); ar(p3,n)

    p2=doc.add_paragraph(); fp(p2); ar(p2,"Summary of Scores:",bold=True)

    scols=[8.922,1.429,2.064,2.381,contentW-8.922-1.429-2.064-2.381]
    st=doc.add_table(rows=6,cols=5); st.style="Table Grid"
    for row in st.rows:
        for ci3,cell in enumerate(row.cells):
            cell.width=Cm(scols[ci3])
            for para in cell.paragraphs: fp(para)
    for ci3,h in enumerate(["Key Areas","Score","Average","Sub-Rating","Description"]):
        p3=st.rows[0].cells[ci3].paragraphs[0]; fp(p3,WD_ALIGN_PARAGRAPH.CENTER); ar(p3,h,bold=True)
    sdata=[
        ("Instructional Competence (40%)",info.get("r1",{}).get("score",""),info.get("r1",{}).get("avg",""),info.get("r1",{}).get("sub",""),info.get("r1",{}).get("desc","")),
        ("Learning Environment (20%)",info.get("r2",{}).get("score",""),info.get("r2",{}).get("avg",""),info.get("r2",{}).get("sub",""),info.get("r2",{}).get("desc","")),
        ("Interpersonal Relations (20%)",info.get("r3",{}).get("score",""),info.get("r3",{}).get("avg",""),info.get("r3",{}).get("sub",""),info.get("r3",{}).get("desc","")),
        ("Professionalism & Personal Characteristics (20%)",info.get("r4",{}).get("score",""),info.get("r4",{}).get("avg",""),info.get("r4",{}).get("sub",""),info.get("r4",{}).get("desc","")),
        ("Total","","",info.get("total",""),info.get("totalDesc",""))
    ]
    for ri3,rd in enumerate(sdata):
        row=st.rows[ri3+1]
        for ci3,val in enumerate(rd):
            p3=row.cells[ci3].paragraphs[0]
            al=WD_ALIGN_PARAGRAPH.RIGHT if (ri3==4 and ci3==0) else (WD_ALIGN_PARAGRAPH.CENTER if ci3>0 else WD_ALIGN_PARAGRAPH.LEFT)
            fp(p3,al); ar(p3,str(val),bold=(ri3==4))

    p2=doc.add_paragraph(); fp(p2)
    ar(p2,"Rated by: ",bold=True); ar(p2,info.get("rater","")+" ______________________")
    ar(p2,"\t\t\t\t\t"); ar(p2,"Date: ",bold=True); ar(p2,info.get("ratedate","_____________________"))
    p2=doc.add_paragraph(); fp(p2); ar(p2,"                              "+info.get("ptitle","School Head"))
    p2=doc.add_paragraph(); fp(p2)
    ar(p2,"Copy received by: ",bold=True); ar(p2,info.get("recv","")+" ________________________________")
    ar(p2,"\t\t"); ar(p2,"Date: ",bold=True); ar(p2,"_____________________")
    p2=doc.add_paragraph(); fp(p2); ar(p2,"                   \tName & Signature of Teacher")

    return doc

data=json.load(sys.stdin)
doc=make_doc(data)
buf=io.BytesIO(); doc.save(buf); buf.seek(0)
print(base64.b64encode(buf.read()).decode())
