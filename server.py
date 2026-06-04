from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io, os, re, webbrowser, threading, time

app = Flask(__name__, static_folder='.', static_url_path='')
CORS(app)

def make_doc(data):
    ratings = data.get("ratings", {})
    info    = data.get("info", {})
    template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "template.docx")
    doc = Document(template_path)

    def fill_para(para, value_map):
        pos = 0
        for run in para.runs:
            if "_" in run.text:
                if pos in value_map:
                    v = str(value_map[pos])
                    m = re.search(r"_+", run.text)
                    if m:
                        ul_len = len(m.group())
                        run.text = re.sub(r"_+", v.ljust(ul_len,"_") if len(v)<=ul_len else v, run.text, count=1)
                pos += 1

    def keep_together(para):
        pPr = para._p.get_or_add_pPr()
        for tag in ["w:keepLines","w:keepNext"]:
            for ex in pPr.findall(qn(tag)): pPr.remove(ex)
        pPr.append(OxmlElement("w:keepLines"))
        pPr.append(OxmlElement("w:keepNext"))

    for pi in [19,20,21,23,24,25,26]:
        try: keep_together(doc.paragraphs[pi])
        except: pass

    for row in doc.tables[5].rows:
        trPr = row._tr.get_or_add_trPr()
        for cs in trPr.findall(qn("w:cantSplit")): trPr.remove(cs)
        cs = OxmlElement("w:cantSplit")
        cs.set(qn("w:val"), "1")
        trPr.append(cs)

    name=info.get("name",""); cls=info.get("cls","")
    subj=info.get("subj",""); dobs=info.get("dateobs","")
    rater=info.get("rater",""); ptitle=info.get("ptitle","School Head")
    rdate=info.get("ratedate",""); recv=info.get("recv","")
    r1=info.get("r1",{}); r2=info.get("r2",{}); r3=info.get("r3",{}); r4=info.get("r4",{})
    total=info.get("total",""); tdesc=info.get("totalDesc","")

    fill_para(doc.paragraphs[5],  {0:name, 1:cls})
    fill_para(doc.paragraphs[7],  {0:subj, 1:dobs})
    for pi,r in [(11,r1),(13,r2),(15,r3),(18,r4)]:
        fill_para(doc.paragraphs[pi], {0:str(r.get("score","")),1:str(r.get("avg","")),2:str(r.get("desc","")),3:str(r.get("sub",""))})
    fill_para(doc.paragraphs[23], {0:rater, 1:rdate})
    for run in doc.paragraphs[24].runs:
        run.text = run.text.replace("School Head", ptitle)
    fill_para(doc.paragraphs[25], {0:recv})

    for sid,ti in [("s1",0),("s2",1),("s3",2),("s4",3)]:
        tbl=doc.tables[ti]; ri2=0
        for row_idx in range(1,len(tbl.rows)):
            row=tbl.rows[row_idx]
            if "Demonstrates" in row.cells[0].text: continue
            sel=int(ratings.get(sid,{}).get(str(ri2),0))
            for ci,v in enumerate([5,4,3,2,1]):
                cell=row.cells[ci+1]
                p_el=cell.paragraphs[0]._p
                for r in p_el.findall(qn("w:r")): p_el.remove(r)
                run=cell.paragraphs[0].add_run("\u2713" if sel==v else "")
                run.font.name="Calibri"; run.font.size=Pt(11)
            ri2+=1

    st=doc.tables[5]
    sd=[(r1.get("score",""),r1.get("avg",""),r1.get("sub",""),r1.get("desc","")),
        (r2.get("score",""),r2.get("avg",""),r2.get("sub",""),r2.get("desc","")),
        (r3.get("score",""),r3.get("avg",""),r3.get("sub",""),r3.get("desc","")),
        (r4.get("score",""),r4.get("avg",""),r4.get("sub",""),r4.get("desc","")),
        ("","",total,tdesc)]
    for ri3,rd in enumerate(sd):
        row=st.rows[ri3+1]
        for ci3,val in enumerate(rd):
            cell=row.cells[ci3+1]
            p_el=cell.paragraphs[0]._p
            for r in p_el.findall(qn("w:r")): p_el.remove(r)
            run=cell.paragraphs[0].add_run(str(val))
            run.font.name="Calibri"; run.font.size=Pt(11)
            if ri3==4: run.bold=True

    return doc

@app.route("/")
def index():
    return send_from_directory(".", "index.html")

@app.route("/<path:filename>")
def static_files(filename):
    return send_from_directory(".", filename)

@app.route("/generate-docx", methods=["POST","OPTIONS"])
def generate_docx():
    if request.method == "OPTIONS":
        return "", 204
    try:
        data = request.get_json()
        doc = make_doc(data)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True, download_name="GLA_Teacher_Performance_Rating.docx")
    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "trace": traceback.format_exc()}), 500

@app.route("/health")
def health():
    return jsonify({"status": "ok"})

def open_browser():
    time.sleep(1.5)
    webbrowser.open("http://localhost:5050")

if __name__ == "__main__":
    print("=" * 50)
    print("  GLA Teacher Rating System")
    print("  Opening http://localhost:5050")
    print("  Keep this window open.")
    print("=" * 50)
    threading.Thread(target=open_browser, daemon=True).start()
    app.run(host="0.0.0.0", port=5050, debug=False)
